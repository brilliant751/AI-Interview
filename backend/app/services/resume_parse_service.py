"""简历解析服务。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


class ResumeParseService:
    """负责将简历文件解析为可用于提问的纯文本。"""

    def parse(self, filename: str, content: bytes) -> str:
        """按扩展名分发解析逻辑。"""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(content)
        if ext == ".docx":
            return self._parse_docx(content)
        if ext == ".doc":
            return self._parse_doc(content)
        raise ValueError("不支持的简历格式")

    def _parse_pdf(self, content: bytes) -> str:
        """解析 PDF 文本。"""
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        parts: list[str] = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
        return self._normalize_text("\n".join(parts))

    def _parse_docx(self, content: bytes) -> str:
        """解析 DOCX 文本（段落 + 表格）。"""
        from docx import Document

        document = Document(BytesIO(content))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return self._normalize_text("\n".join(parts))

    def _parse_doc(self, content: bytes) -> str:
        """解析老式 DOC 文本（受限降级）。"""
        # 说明：不引入系统命令依赖时，.doc 的稳定解析能力有限，这里做最佳努力提取。
        for encoding in ("utf-8", "gb18030", "latin-1"):
            try:
                text = content.decode(encoding, errors="ignore")
            except Exception:
                continue
            cleaned = self._normalize_text(text)
            if len(cleaned) >= 80:
                return cleaned
        return ""

    def _normalize_text(self, text: str) -> str:
        """统一文本清洗并限制长度。"""
        rows = [row.strip() for row in text.replace("\x00", "").splitlines()]
        merged = "\n".join(row for row in rows if row)
        return merged[:12000]
